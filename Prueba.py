# import pandas as pd
# from docx import Document
# from docx.shared import Pt
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from docx.enum.table import WD_ALIGN_VERTICAL
# from docx.shared import Inches

# # Cargar el archivo Excel
# archivo_excel = 'C:\\Users\\Desarrollo\\Desktop\\Defensa Civil Proyecto 2022.xlsx'
# datos_excel = pd.read_excel(archivo_excel)

# # Crear un nuevo documento Word con orientación vertical
# documento_word = Document()
# section = documento_word.sections[0]
# section.start_type
# new_width, new_height = section.page_height, section.page_width
# section.page_width = new_width
# section.page_height = new_height

# # Iterar a través de los registros en el DataFrame y agregar una hoja por registro al documento Word
# for indice, fila in datos_excel.iterrows():
#     # Crear una nueva página en el documento Word
#     documento_word.add_page_break()
    
#     # Agregar un encabezado con el nombre del registro
#     documento_word.add_heading(f'Registro {indice + 1}', level=1)
    
#     # Agregar una nueva tabla al documento Word
#     tabla = documento_word.add_table(rows=len(datos_excel.columns) + 2, cols=2)
    
#     # Ajustar el ancho de las celdas
#     for row in tabla.rows:
#         for cell in row.cells:
#             cell.width = Pt(100)  # Ajustar el ancho según sea necesario
    
#     # Agregar encabezados (incluyendo el índice)
#     encabezados = tabla.rows[0].cells
#     encabezados[0].text = "Índice"
#     encabezados[1].text = "Valor"
    
#     # Agregar datos de la fila al documento Word
#     for i, columna in enumerate(datos_excel.columns, start=1):
#         nueva_fila = tabla.rows[i].cells
#         nueva_fila[0].text = columna
#         nueva_fila[1].text = str(fila[columna])

#     # Agregar una fila adicional para la imagen
#     imagen_fila = tabla.rows[len(datos_excel.columns) + 1].cells
    
#     # Insertar la imagen (ajusta la ruta a la ubicación real de tus imágenes)
#     ruta_imagen = 'ruta/de/imagen.png'
#     imagen_fila[0].text = 'Imagen'
#     imagen_fila[1].paragraphs[0].add_run().add_picture(ruta_imagen, width=Inches(2.0))

# # Guardar el documento Word
# documento_word.save('C:\\Users\\Desarrollo\\Desktop\\Defensa Civil Proyecto 2022.docx')











import pandas as pd
from docx import Document
from docx.shared import Pt

# Cargar el archivo Excel
archivo_excel = 'C:\\Users\\Desarrollo\\Desktop\\Defensa Civil Proyecto 2022.xlsx'
datos_excel = pd.read_excel(archivo_excel)

# Crear un nuevo documento Word con orientación vertical
documento_word = Document()
section = documento_word.sections[0]
section.start_type
new_width, new_height = section.page_height, section.page_width
section.page_width = new_width
section.page_height = new_height

# Iterar a través de los registros en el DataFrame y agregar una hoja por registro al documento Word
for indice, fila in datos_excel.iterrows():
    # Crear una nueva página en el documento Word
    documento_word.add_page_break()
    
    # Agregar un encabezado con el nombre del registro
    documento_word.add_heading(f'Registro {indice + 1}', level=1)
    
    # Agregar una nueva tabla al documento Word
    tabla = documento_word.add_table(rows=len(datos_excel.columns) + 1, cols=2)
    
    # Ajustar el ancho de las celdas
    for row in tabla.rows:
        for cell in row.cells:
            cell.width = Pt(100)  # Ajustar el ancho según sea necesario
    
    # Agregar encabezados (incluyendo el índice)
    encabezados = tabla.rows[0].cells
    encabezados[0].text = "Índice"
    encabezados[1].text = "Valor"
    
    # Agregar datos de la fila al documento Word
    for i, columna in enumerate(datos_excel.columns, start=1):
        nueva_fila = tabla.rows[i].cells
        nueva_fila[0].text = columna
        nueva_fila[1].text = str(fila[columna])

# Guardar el documento Word
documento_word.save('C:\\Users\\Desarrollo\\Desktop\\Defensa Civil Proyecto 2022.docx')
print('Todo ok')